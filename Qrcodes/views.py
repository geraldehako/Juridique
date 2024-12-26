import qrcode
from io import BytesIO
from django.core.files.base import ContentFile
from docx import Document
from django.http import HttpResponse
from django.shortcuts import render, redirect, get_object_or_404
from .models import Attestation
from .forms import ClientForm, AttestationForm
from docx.shared import Inches

def creer_attestation(request):
    if request.method == 'POST':
        # Récupération des données POST
        client_form = ClientForm(request.POST)
        attestation_form = AttestationForm(request.POST)

        # Vérification de la validité des formulaires
        if client_form.is_valid() and attestation_form.is_valid():
            # Sauvegarde du client et attestation
            client = client_form.save()  # Sauvegarde du client
            attestation = attestation_form.save(commit=False)  # L'attestation sans sauvegarde immédiate
            attestation.client = client  # Lier l'attestation au client
            attestation.save()  # Sauvegarde de l'attestation

            # Rediriger vers la page de téléchargement
            return redirect('telecharger_attestation', attestation_id=attestation.id)
        else:
            # En cas d'erreur dans les formulaires, afficher les erreurs
            return render(request, 'Pages/qr/creer_attestation.html', {
                'client_form': client_form,
                'attestation_form': attestation_form,
                'errors': client_form.errors if not client_form.is_valid() else attestation_form.errors
            })
    else:
        # Initialiser les formulaires vides si méthode GET
        client_form = ClientForm()
        attestation_form = AttestationForm()

    return render(request, 'Pages/qr/creer_attestation.html', {
        'client_form': client_form,
        'attestation_form': attestation_form,
    })

def liste_attestations(request):
    # Affiche la liste de toutes les attestations
    attestations = Attestation.objects.all()
    return render(request, 'Pages/qr/liste_attestations.html', {'attestations': attestations})

def generer_qr_code(attestation):
    # Création du QR code avec les informations de l'attestation
    qr_data = (
        f"Type Attestion: {attestation.type_attestation}\n"
        f"Numéro: {attestation.numero}\n"
        f"Client: {attestation.client.nom} {attestation.client.prenoms}\n"
        f"CNI: {attestation.client.numero_cni}\n"
        f"Lotissement: {attestation.lotissement}\n"
        f"Lot Numéro: {attestation.lot_numero}\n"
        f"Superficie: {attestation.superficie} m²\n"
        f"Arrêté Référence: {attestation.arrete_reference}"
    )

    # Utilisation de la bibliothèque qrcode pour créer l'image
    qr = qrcode.QRCode(version=1, box_size=10, border=4)
    qr.add_data(qr_data)
    qr.make(fit=True)

    img = qr.make_image(fill='black', back_color='white')
    buffer = BytesIO()
    img.save(buffer, format='PNG')  # Sauvegarde de l'image dans le buffer
    return buffer.getvalue()

from docx import Document
from io import BytesIO
from django.http import HttpResponse
from docx.shared import Inches

def generer_document_word(attestation, qr_code):
    # Création du document Word
    doc = Document()

    # Ajout du titre
    doc.add_heading("ATTESTATION DE {attestation.type_attestation.capitalize()}", level=1)
    doc.add_paragraph(f"Numéro : {attestation.numero}", style="Intense Quote")

    # Bloc d'informations principales
    doc.add_paragraph(
        "Nous soussignés, NANAN KOUAME YAO, Chef du village de N'GBEKRO, sous-préfecture de Lolobo, "
        "et KOUADIO KOUAKOU ISIDORE, Président de la Commission Foncière et des Lotissements, attestons que :"
    )

    # Détails du client
    doc.add_paragraph(f"Mme/Mlle/Mr : {attestation.client.nom} {attestation.client.prenoms}")
    doc.add_paragraph(f"Né(e) le : {attestation.client.date_naissance}")
    doc.add_paragraph(f"Adresse : {attestation.client.adresse}")

    # Informations sur le lotissement
    doc.add_paragraph(
        f"Est incontestablement et de façon irrévocable attributaire du "
        f"Lot N°{attestation.lot_numero} d'une superficie de {attestation.superficie} m² issu du lotissement dénommé "
        f"{attestation.lotissement}, approuvé par arrêté N°{attestation.arrete_reference}."
    )

    # Mention de cession
    doc.add_paragraph(
        f"A ce titre, jouissant de son plein droit, il/elle cède ledit lot à :"
    )
    #doc.add_paragraph(f"Nom de l'acquéreur : {attestation.acquereur.nom_complet}")
    #doc.add_paragraph(f"Adresse de l'acquéreur : {attestation.acquereur.adresse}")
    #doc.add_paragraph(f"Document d'identité : {attestation.acquereur.numero_identite}")
    doc.add_paragraph(f"Nom de l'acquéreur : {attestation.client.prenoms}")
    doc.add_paragraph(f"Adresse de l'acquéreur : {attestation.client.prenoms}")
    doc.add_paragraph(f"Document d'identité : {attestation.client.prenoms}")

    # Date et signature
    doc.add_paragraph(
        f"Fait à N'GBEKRO, le {attestation.client.date_naissance.strftime('%d/%m/%Y')}."
        # f"Fait à N'GBEKRO, le {attestation.date_creation.strftime('%d/%m/%Y')}."
    )
    doc.add_paragraph("\n")
    doc.add_paragraph("Le Président de la Commission Foncière et des Lotissements")
    doc.add_paragraph("KOUADIO KOUAKOU ISIDORE", style="Heading 2")

    # Insertion du QR code
    qr_paragraph = doc.add_paragraph("\n")
    run = qr_paragraph.add_run()
    run.add_picture(BytesIO(qr_code), width=Inches(1.5))

    # Enregistrement du fichier Word dans la réponse
    file_name = f"attestation_{attestation.numero}.docx"
    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response["Content-Disposition"] = f'attachment; filename="{file_name}"'
    doc.save(response)
    return response


def generer_document_word1(attestation, qr_code):
    # Création du document Word
    doc = Document()

    # Ajout des données de l'attestation dans le document
    doc.add_heading(f"Attestation {attestation.type_attestation.capitalize()}", level=1)
    doc.add_paragraph(f"Numéro : {attestation.numero}")
    doc.add_paragraph(f"Nom du client : {attestation.client.nom} {attestation.client.prenoms}")
    doc.add_paragraph(f"Date de naissance : {attestation.client.date_naissance}")
    doc.add_paragraph(f"Adresse : {attestation.client.adresse}")
    doc.add_paragraph(f"Numéro CNI : {attestation.client.numero_cni}")
    doc.add_paragraph(f"Lotissement : {attestation.lotissement}")
    doc.add_paragraph(f"Lot Numéro : {attestation.lot_numero}")
    doc.add_paragraph(f"Superficie : {attestation.superficie} m²")
    doc.add_paragraph(f"Arrêté Référence : {attestation.arrete_reference}")

    # Insertion du QR code dans le document
    qr_paragraph = doc.add_paragraph()
    run = qr_paragraph.add_run()
    run.add_picture(BytesIO(qr_code), width=Inches(1.5))

    # Enregistrement du fichier Word dans la réponse
    file_name = f"attestation_{attestation.numero}.docx"
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{file_name}"'
    doc.save(response)
    return response

def telecharger_attestation(request, attestation_id):
    # Récupérer l'attestation à partir de l'ID
    attestation = get_object_or_404(Attestation, id=attestation_id)

    # Générer le QR code et le document Word
    qr_code = generer_qr_code(attestation)
    return generer_document_word(attestation, qr_code)

from django.shortcuts import render
from .models import QR_code

def index(request, *args, **kwargs):
    context = {}

    if request.method == "POST":
        data = request.POST.get('data')
        # Création de l'objet QR_code dans la base de données
        obj = QR_code.objects.create(data=data)
        qr_code = obj.qr_code.url  # Accès à l'URL de l'image générée

        # Ajout de l'image au contexte pour l'afficher dans le template
        context = {
            'qr_code': qr_code
        }

    return render(request, 'Pages/qr/index.html', context)

